#!/usr/bin/env python
"""Configuration Tool

The Configuration Tool is a software program produced for the European Space
Agency.

The purpose of this tool is to keep under configuration control the changes
in the Ground Segment components of the Copernicus Programme, in the
framework of the Coordination Desk Programme, managed by Telespazio S.p.A.

This program is free software: you can redistribute it and/or modify it under
the terms of the GNU General Public License as published by the Free Software
Foundation, either version 3 of the License, or (at your option) any later
version.

This program is distributed in the hope that it will be useful, but WITHOUT
ANY WARRANTY; without even the implied warranty of MERCHANTABILITY or FITNESS
FOR A PARTICULAR PURPOSE. See the GNU General Public License for more details.
You should have received a copy of the GNU General Public License along with
this program. If not, see <http://www.gnu.org/licenses/>.
"""

__author__ = "Coordination Desk Development Team"
__contact__ = "coordination_desk@telespazio.com"
__copyright__ = "Copyright 2024, Telespazio S.p.A."
__license__ = "GPLv3"
__status__ = "Production"
__version__ = "1.0.0"

import flask
import json
import logging
import io
import re
import sys
import tempfile
import threading

from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime as dt
from enum import Enum
from functools import cmp_to_key
from pathlib import Path
from typing import (
    Callable, ClassVar, Dict, Iterable, Iterator, List, Optional, Tuple,
    TypedDict, Union
)

if sys.version_info >= (3, 10):
    from typing import TypeAlias
else:
    from typing_extensions import TypeAlias

import docx
import docx.document
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tc
from docx.shared import Inches
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from html4docx import HtmlToDocx

from pathlib import Path
if str(Path(__file__).parent.parent.parent) not in sys.path:
    sys.path.append(str(Path(__file__).parent.parent.parent))

from apps.utils import file_utils


CURRENT_DATAFLOW_DOC_VERSION = "1.8"
CURRENT_DATAFLOW_DOC_TEMPLATE = "apps/config/templates/dataflow/[EOF-DFC] EOF_CSC Sentinels Data Flow Configuration - Empty Template.docx"

# ########################################################################### #
# ############################ Enums ######################################## #
# ########################################################################### #
# region Enums
class Mission(Enum):
    """
    All possible Sentinel missions. Values are as found in the doc table and
    names are as found in the DB
    """
    S1 = "Sentinel-1"
    S2 = "Sentinel-2"
    S3 = "Sentinel-3"
    S5P = "Sentinel-5P"
    S6 = "Sentinel-6"

class Group(Enum):
    """
    Enumeration whose values are readable group names as used in dataflow doc
    """
    PRODUCTS = "Products"
    AUX_DATA = "Auxiliary Data"
    MP_AND_FOS = "Mission Planning and FOS Files"
    OLQC_REPORTS = "OLQC Reports"
    REMOVED_PRODUCTS = "Removed Products"
    REMOVED_AUX_DATA = "Removed Auxiliary Data"
    REMOVED_MP_AND_FOS = "Removed Mission Planning and FOS Files"

# Tuple of all groups containing 'Removed' products
REMOVED_GROUPS = (Group.REMOVED_AUX_DATA, Group.REMOVED_MP_AND_FOS, Group.REMOVED_PRODUCTS)

class Column(Enum):
    """
    Enumeration whose values are column header values in the dataflow document's tables
    """
    DESCRIPTION = "Description"
    PRODUCT_TYPE = "Product Type"
    PAYLOAD = "Payload"
    LEVEL = "Level"
    MODE = "Mode"
    TYPE = "Type"

class Service(Enum):
    """
    Enumeration whose values are service values in the dataflow JSON data
    
    Notes
    -----
    Upon initialisation via 'Service(<consumer/producer string>)', if the
    provided string is not a known value, Service.Unknown is instantiated
    """
    ADG = "ADG"          # Auxiliary Data Group
    DA = "DA"            # Data Access Service
    E2E = "E2E"          # End-to-End Service
    EDRS = "EDRS"        # European Data Relay Satellite
    EUM = "EUM"          # EUMETSAT
    EXT = "EXT"          # External Provider
    FOS = "FOS"          # Flight Operations Service
    LTA = "LTA"          # Long-Term Archive
    MP = "MP"            # Mission Planning
    MPC = "MPC"          # Mission Performance Cluster
    POD = "POD"          # Precise Orbit Determination
    PR = "PR"            # Production
    RS = "RS"            # Reference System
    X_BAND = "X-Band"    # X-Band Station
    UNKNOWN = "UNKNOWN"  # For unmapped services in the JSON data, to avoid
                         # parsing errors and allow for manual review of 
                         # unmapped entries

    @classmethod
    def _missing_(cls, *args, **kwargs):
        return cls.UNKNOWN

class Entity(Enum):
    """
    Enum of all possible producers/consumers whose values are as written in
    the dataflow doc tables.
    
    Notes
    -----
    Upon initialisation via 'Entity(<consumer/producer string>)', if the
    provided string is not a known value, Entity.Unknown is instantiated
    """
    C = "C"                  # Consumer
    C_ADG = "C-ADG"          # Consumer of Auxiliary Data Group
    C_DA = "C-DA"            # Consumer of Data Access Service
    C_EDRS = "C-EDRS"        # Consumer of European Data Relay Satellite
    C_EUM = "C-EUM"          # Consumer of EUMETSAT
    C_FOS = "C-FOS"          # Consumer of Flight Operations Service
    C_LTA = "C-LTA"          # Consumer of Long-Term Archive
    C_MP = "C-MP"            # Consumer of Mission Planning
    C_MPC = "C-MPC"          # Consumer of Mission Performance Cluster
    C_NOAA = 'C-NOAA'        # Consumer of [?]
    C_NPP = "C-NPP"          # Consumer of Suomi National Polar-orbiting Partnership
    C_OSI_SAF = "C-OSI_SAF"  # Consumer of [?]
    C_POD = "C-POD"          # Consumer of Precise Orbit Determination
    C_PR = "C-PR"            # Consumer of Production
    C_RS = "C-RS"            # Consumer of Reference System
    P = "P"                  # Provider
    P0 = "P0"                # Original Provider
    UNKNOWN = "UNKNOWN"      # For unmapped services in the JSON data, to avoid
                             # parsing errors and allow for manual review of 
                             # unmapped entries
    @classmethod
    def _missing_(cls, *args, **kwargs):
        print("UNKNOWN: ", *args)
        return cls.UNKNOWN

CONSUMER_ENTITIES = tuple(e for e in Entity if e.value.startswith('C-'))
PRODUCER_ENTITIES = tuple(e for e in Entity if e.value.startswith('P'))
#endregion
# ########################################################################### #
# ############################ Classes ###################################### #
# ########################################################################### #
# region classes
class SectionContent(TypedDict):
    """Type for WordGenerator.doc_structure dict's values"""
    number: str  # Section number as a string e.g. 1.2.3
    tables: list[docx.table.Table]  # List of tables

def _track_changes_wrapper(func):
    """
    Decorator for WordGenerator methods: skips execution if WordGenerator's
    tracking_changes attribute is False.
    """
    from functools import wraps

    @wraps(func)
    def wrapper(self, *args, **kwargs):
        if not self.tracking_changes:
            return
        return func(self, *args, **kwargs)
    return wrapper

class WordGenerator:
    """
    Wrapper class for docx.document.Document
    
    Parameters
    ----------
    path_to_document: Optional[Union[str, Path]]
        System path to the existing document, if needed

    Attributes
    ----------
    doc_structure: dict[Tuple[str, ...], SectionContent]
        A dictionary whose keys are tuples of section/sub-section headers and
        values are a dict whose keys are 'number' and 'tables', whose values 
        are a str representing section number e.g. 3.4.1, and a list of 
        docx.table.Table instances of tables in that section, respectively

    tracking_changes: bool
        Whether to track changes or not. Default is True. If False, 
        make_ins_run, insert_existing_runs, and delete_existing_runs methods
        do nothing

    __document: docx.document.Document
        Wrapped docx.document.Document instance

    __map: dict
        Dictionary mapping to track added paragraphs, pictures, tables, and 
        headings. TODO: Fold into doc_structure attribute

    _current_track_change_id: int
        The id of the latest track change
    """
    def __init__(self, path_to_document: Optional[Union[str, Path]] = None):
        self.doc_structure: dict[Tuple[str, ...], SectionContent] = {}
        self.tracking_changes: bool = True

        self.__document: docx.document.Document = None
        self.__map = {}
        self.create(path_to_document)
        self._get_doc_structure()

        self._current_track_change_id = self._get_most_recent_track_change_id()

    def create(self, path_to_document=None):
        """
        :return:
        :rtype:
        """
        if path_to_document:
            self.__document = Document(path_to_document)
        else:
            self.__document = Document()

    def add_title_header(self, title):
        """
        :param title:
        :type title:
        :return:
        :rtype:
        """
        section = self.__document.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = title
        return header

    def add_heading(self, name, level):
        """
        :param name:
        :type name:
        :param level:
        :type level:
        :return:
        :rtype:
        """
        self.__map[name] = self.__document.add_heading(name, level)
        return self.__map[name]

    def add_paragraph(self, name, text):
        """
        :param name:
        :type name:
        :param text:
        :type text:
        :return:
        :rtype:
        """
        self.__map[name] = self.__document.add_paragraph(text)
        return self.__map[name]

    def add_paragraph_before(self, paragraph, text=None, style=None):
        for par in self.__document.paragraphs:
            if par.text.lower() == paragraph.text.lower():
                par.insert_paragraph_before(text, style)
        return None

    def add_paragraph_after(self, paragraph, text=None, style=None):
        for index, par in enumerate(self.__document.paragraphs):
            if par.text.lower() == paragraph.text.lower():
                next_par = self.__document.paragraphs[index + 1]
                added_par = next_par.insert_paragraph_before(text, style)
                return added_par
        return None

    def add_picture(self, image_path, paragraph=None, caption=None):
        """
        :param paragraph:
        :param image_data:
        :type image_data:
        :return:
        :rtype:
        """
        # skip function if image is null
        if image_path is None:
            return

        # adjust the size as needed
        width, height = 5, 4.5

        # Use a unique identifier for each image
        image_id = len(self.__map) + 1
        if paragraph is None:
            self.__map[image_id] = self.__document.add_picture(
                image_path, width=Inches(width), height=Inches(height)
            )
            last_paragraph = self.__document.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            new_p = OxmlElement("w:p")
            paragraph._p.addnext(new_p)
            added_par = Paragraph(new_p, paragraph._parent)
            run = added_par.add_run()
            self.__map[image_id] = run.add_picture(
                image_path, width=Inches(width), height=Inches(height)
            )
            added_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # If present, add a figure caption
            if caption:
                # New caption paragraph
                new_p = OxmlElement("w:p")
                added_par._p.addnext(new_p)
                caption_par = Paragraph(new_p, added_par._parent)
                caption_par.style = "Caption"
                caption_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Add figure label
                caption_par.add_run("Figure ")

                # Add automatic numbering
                run = caption_par.add_run()
                r = run._r

                fldChar = docx.oxml.OxmlElement("w:fldChar")
                fldChar.set(docx.oxml.ns.qn("w:fldCharType"), "begin")
                r.append(fldChar)

                instrText = docx.oxml.OxmlElement("w:instrText")
                instrText.text = " SEQ Figure \\* ARABIC"
                r.append(instrText)

                fldChar = docx.oxml.OxmlElement("w:fldChar")
                fldChar.set(docx.oxml.ns.qn("w:fldCharType"), "end")
                r.append(fldChar)

                # Add text
                caption_par.add_run(" " + caption)

        return image_id

    def add_table(self, name, rows, cols, paragraph=None):
        """
        :param paragraph:
        :param name:
        :type name:
        :param rows:
        :type rows:
        :param cols:
        :type cols:
        :return:
        :rtype:
        """
        self.__map[name] = self.__document.add_table(rows=rows, cols=cols)
        if paragraph is not None:
            paragraph._p.addnext(self.__map[name]._tbl)
        return self.__map[name]

    def add_text_to_cell_table(self, table_name, row, col, text):
        """
        :param table_name:
        :type table_name:
        :param row:
        :type row:
        :param col:
        :type col:
        :param text:
        :type text:
        :return:
        :rtype:
        """
        cell = self.__map[table_name].cell(row, col)
        parser = HtmlToDocx()
        parser.add_html_to_cell(text, cell)
        return cell

    def set_horizontal_layout(self):
        sections = self.__document.sections
        for section in sections:
            new_width, new_height = section.page_height, section.page_width
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height

    def get_paragraph(self, paragraph_name):
        for index, par in enumerate(self.__document.paragraphs):
            if par.text.lower() == paragraph_name.lower():
                return par
        return None

    def set_section_width_height(self, index, width, height):
        section = self.__document.sections[index]
        section.page_width = width
        section.page_height = height

    def set_vertical_cell_direction(self, cell: _Cell, direction: str):
        # direction: tbRl -- top to bottom, btLr -- bottom to top
        assert direction in ("tbRl", "btLr")
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        textDirection = OxmlElement("w:textDirection")
        textDirection.set(qn("w:val"), direction)  # btLr tbRl
        tcPr.append(textDirection)

    def save(self, name):
        """
        :param name:
        :type name:
        :return:
        :rtype:
        """
        path = tempfile.gettempdir()
        self.__document.save(
            path + "/" + name + " - " + file_utils.get_date_for_file() + ".docx"
        )
        return path + "/" + name + " - " + file_utils.get_date_for_file() + ".docx"

    @property
    def document(self) -> Document:
        return self.__document

    def get_mission_group_table(self, mission: Mission, group: Group) -> Tuple[docx.table.Table, ...]:
        """"""
        for section in self.doc_structure:
            section_str = ' '.join(section).lower()

            if mission.value.lower() not in section_str:
                continue

            if group.value.lower() not in section_str:
                continue

            return tuple(self.doc_structure[section]['tables'])
        return ()

    def _get_most_recent_track_change_id(self) -> int:
        """
        Get the most recent id attribute value associated with tracked changes
        in the document. Returns -1 if no track changes are present
        """
        body = self.document.element.body
        track_changes_ids = {
            int(el.get(qn('w:id')))
            for el in body.iter()
            if el.tag in (qn('w:ins'), qn('w:del'))
            and el.get(qn('w:id')) is not None
        }

        return max(track_changes_ids) if len(track_changes_ids) > 0 else -1

    def _get_doc_structure(self):
        """
        Evaluates current document structure, and writes the results to 
        self.doc_structure
        """
        body = self.document.element.body
        current_level = 0
        current_nesting = []
        current_section = []
        is_numbered_section = False
        for child in body:
            is_table = child.tag == qn("w:tbl")
            if is_table:
                try:
                    table_idx = [_._element == child for _ in self.document.tables].index(True)
                    self.doc_structure[tuple(current_nesting)]['tables'].append(self.document.tables[table_idx])
                except ValueError as err:
                    pass
                continue

            if child.tag == qn("w:p"):
                para = Paragraph(child, self.document)
                heading = re.match('Heading([0-9]+)', para.style.name)  # Numbered sections

                if heading:
                    is_numbered_section = True
                    heading_level = int(heading.group(1))
                else:
                    is_numbered_section = False
                    heading = re.match('DOC_TYPE', para.style.name)  # Unnumbered sections
                    heading_level = None

                if heading:
                    section_name = para.text

                    if heading_level is not None:
                        if current_level is None or heading_level > current_level:
                            if current_level is None:
                                current_nesting = []
                            current_nesting.append(section_name)
                            current_section.append(1)
                        else:
                            if is_numbered_section:
                                if current_section:
                                    current_section = current_section[:heading_level - 1] + [current_section[heading_level - 1] + 1]
                                else:
                                    current_section = current_section[:heading_level - 1]
                                current_nesting = current_nesting[:heading_level - 1] + [section_name]
                            else:
                                current_nesting = [section_name]
                    else:
                        current_nesting = [section_name]

                    self.doc_structure[tuple(current_nesting)] = {
                        'number': '.'.join([str(_) for _ in current_section]) if is_numbered_section else '',
                        'tables': [],
                    }
                    current_level = heading_level

    @_track_changes_wrapper
    def make_ins_run(self, paragraph, text, author=__author__):
        """
        Adds a run wrapped in <w:ins> so it appears as a tracked insertion, if
        tracking_changes is set to True
        """
        self._current_track_change_id += 1
        ins = OxmlElement('w:ins')
        ins.set(qn('w:id'), str(self._current_track_change_id))
        ins.set(qn('w:author'), author)
        ins.set(qn('w:date'), dt.now().strftime("%Y-%m-%dT%H:%M:%S"))

        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        run.append(rPr)

        t = OxmlElement('w:t')
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        run.append(t)

        ins.append(run)
        paragraph._p.append(ins)

    @_track_changes_wrapper
    def insert_existing_runs(self, paragraph, author=__author__):
        """
        Marks all existing runs in a paragraph as tracked insertions, if
        tracking_changes is set to True
        """
        date = dt.now().strftime("%Y-%m-%dT%H:%M:%S")
        self._current_track_change_id += 1

        for run in paragraph.runs:
            r_elem = run._r

            ins = OxmlElement('w:ins')
            ins.set(qn('w:id'), str(self._current_track_change_id))
            ins.set(qn('w:author'), author)
            ins.set(qn('w:date'), date)

            r_elem.getparent().replace(r_elem, ins)
            ins.append(r_elem)

    @_track_changes_wrapper
    def delete_existing_runs(self, paragraph, author=__author__):
        """
        Marks all existing runs in a paragraph as tracked deletions, if
        tracking_changes is set to True
        """
        date_str = dt.now().strftime("%Y-%m-%dT%H:%M:%S")
        self._current_track_change_id += 1

        for run in paragraph.runs:
            r_elem = run._r

            # Build the <w:del> wrapper
            delete = OxmlElement('w:del')
            delete.set(qn('w:id'), str(self._current_track_change_id))
            delete.set(qn('w:author'), author)
            delete.set(qn('w:date'), date_str)

            # Replace <w:t> with <w:delText> inside the run
            for t in r_elem.findall(qn('w:t')):
                del_text = OxmlElement('w:delText')
                del_text.text = t.text
                del_text.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                r_elem.replace(t, del_text)

            # Swap the run out for the del-wrapped version
            r_elem.getparent().replace(r_elem, delete)
            delete.append(r_elem)


class DataflowDocTable:
    """
    Class to represent a table in the Dataflow document (only) for a specific mission and group.
    
    Parameters
    ----------
    word_generator: WordGenerator
        WordGenerator instance containing the document

    mission: Mission
        Mission to which this table relates to

    group: Group
        Group to which this table relates to

    Attributes
    ----------
    _word_generator: WordGenerator
        WordGenerator instance containing the document

    _mission: Mission
        Mission to which this table relates to

    _group: Group
        Group to which this table relates to

    _table: docx.table.Table
        docx.table.Table instance of the document's table. Found by searching
        _word_generator for the appropriate table

    _header_row_indices: List[int]
        Table row indices belonging to the table header. Evaluated once upon
        loading of table

    _catalogue: ProductCatalogue
        ProductCatalogue instance containing all table products. Read-only
        access through catalogue property. Initially, lazily-evaluated when
        catalogue property is accessed

    _catalogue_to_table_row_idx_mapping: Dict[CataloguePrimaryKey, List[int]]
        Mapping between product via their ProductCatalogue primary key, and
        their row in the table as List[int] since producers/consumers can be
        split over two rows for a single Product

    __html_parser: HtmlToDocx
        Parser for translating raw HTML to DOCX (XML) format. Initialised in
        constructor for use in add_product_as_row method

    _change_history: List[DataflowDocTable.CellChange]
        List of changes made to the table since initialisation. Relevant to
        future product comparison logic and change logs. Not currently used.

    _data_row_indices: List[int], @property (read-only)
        Row indices of all data (non-header) rows in the table

    catalogue: ProductCatalogue, @property (read-only)
        Public getter for _catalogue

    col_headers: Tuple[str, ...], @property (read-only)
        Tuple of header strings for each unsplit header cell

    columns: Tuple[Union[Column, Service]], @property (read-only)
        Columns of the table as Column/Service Enum instances. Evaluated on
        each call from module's COLS_BY_MISSION_AND_GROUP mapping

    data: Optional[Iterator[dict]], @property (read-only)
        Table data from every row (non-headers only) returned as an iterator
        of dicts whose keys are self.columns, and values are the relevant 
        cell's text values

    n_data_rows: int
        Number of non-header rows in the table
    """
    BODY_FONT_SIZE_PT: float = 9.0
    BOLD_HEADER: bool = True
    HEADER_FONT_SIZE_PT: float = 10.0
    HEADER_BACKGROUND_COLOR = "e6e6e6"
    COLUMN_COLOURS: Dict[Service, str] = {
        Service.ADG: "e5b8b7",
        Service.DA: "d6e3bc",
        # Service.DATA_ACQUISITION: "ccc0d9",
        Service.E2E: "e6e6e6",
        Service.EDRS: "d6e3bc",
        Service.EUM: "8db3e2",
        Service.EXT: "c6d9f1",
        Service.FOS: "bfbfbf",
        Service.LTA: "ccc0d9",
        Service.MP: "ffffcc",
        Service.MPC: "daeef3",
        Service.POD: "f2f2f2",
        Service.PR: "fbd4b4",
        Service.RS: "adadad",
        Service.X_BAND: "e6e6e6",
    }
    
    @dataclass    
    class CellChange:
        column: Union[Column, Service]
        old: str
        new: str

    def __init__(self, word_generator: WordGenerator, mission: Mission, group: Group):
        """
        Loads docx.table.Table and ensures all required headers are present in
        the last header row from the COLS_BY_MISSION_AND_GROUP mapping

        Raises
        ------
        ValueError
            If the correct headers from the COLS_BY_MISSION_AND_GROUP are not
            present
        """
        self._word_generator = word_generator
        self._mission: Mission = mission
        self._group: Group = group

        self._table: docx.table.Table = self._word_generator.get_mission_group_table(mission, group)[0]
        self._header_row_indices: List[int] = []
        self._catalogue: ProductCatalogue = None
        self._catalogue_to_table_row_idx_mapping: Dict[CataloguePrimaryKey, int] = {}

        # Non-implemented attributes. For future data comparison logic
        self._change_history: List[DataflowDocTable.CellChange] = []

        self.__html_parser = HtmlToDocx()

        # Establish which row indices belong to header rows
        last_header_row_encountered = False
        for row_idx, row in enumerate(self._table.rows):
            cells = row.cells
            for cell in cells:
                cell.text = cell.text.strip().strip('\n')

            self._header_row_indices.append(row_idx)
            if all([hdr.text in self.col_headers for hdr in cells]):
                last_header_row_encountered = True
            
            if last_header_row_encountered:
                break

        else:
            raise ValueError(f"Final header row, '{', '.join(self.col_headers)}', not found")

        self._set_table_font_size()

        tbl = self._table._tbl
        rows = tbl.findall(qn('w:tr'))
        for idx in self._header_row_indices:
            tr = rows[idx]
            for tc in tr.findall(qn('w:tc')):
                self._set_cell_alignment(tc, 'center')
                self._set_cell_vertical_alignment(tc, 'center')

        self._set_table_borders(tbl)

    @staticmethod
    def parse_doc_entities_string(
        entities_str: str,
        delimiters: Tuple[str, ...] = ('/', ' ', '\n')
    ) -> Tuple[Entity, ...]:
        """
        Parse the contents of a table cell containing producers/consumers into
        a tuple of Entity instances.
        
        Parameters
        ----------
        entities_str: str
            String from the cell with the producers/consumers. The expected
            format is a list of consumers/producers delimited by any of the
            characters specified in the delimiters arg.

        delimiters: Tuple[str, ...]
            Characters that delimit individual producer/consumer entities from
            one another in the string. Default is '/', ' ', and/or '\n'.

        Returns
        -------
        entities: Tuple[Entity, ...]
            Tuple of Entity instances representing each product/consumer in
            the input string

        Notes
        -----
        Any trailing [0-9] characters on any of the delimited sub-strings are
        ignored as assumed footnote references.
        """
        entities_strs = [entities_str.upper()]
        for delim in delimiters:
            entities_strs = [
                entity_str.strip()
                for string in entities_strs
                for entity_str in string.split(delim)
                if entity_str.strip()
            ]

        # Remove any superscript numbers. P0 is the only Entity that ends with
        # a [0-9] character
        entities = tuple(
            Entity('P0' if e.startswith('P0') else re.sub(r"[0-9]+$", "", e))
            for e in entities_strs
        )

        return entities

    @property
    def catalogue(self) -> 'ProductCatalogue':
        """ProductCatalogue instance of all Products in table"""
        if self._catalogue is None:
            # TODO: Move logic involved below to a separate function
            pk_func = lambda product: (self._mission, self._group, product[Column.PRODUCT_TYPE])
            self._catalogue = ProductCatalogue(pk_func=pk_func)

            service_cols = tuple(col for col in self.columns if isinstance(col, Service))
            for idx, row in enumerate(self.data):
                product = Product(
                    "", self._mission, self._group,
                    row.get(Column.PRODUCT_TYPE, ""),
                    row.get(Column.PAYLOAD, ""),
                    row.get(Column.LEVEL, ""),
                    row.get(Column.MODE, ""),
                    row.get(Column.TYPE, ""),
                    row.get(Column.DESCRIPTION, ""),
                    entities_relations={
                        service: self.parse_doc_entities_string(row[service])
                        for service in service_cols
                    }
                )

                # docx table splits provider/consumer entities on two separate
                # lines. Therefore, need to combine in to each unique product 
                # in catalogue
                catalogue_key = pk_func(product)
                if self._catalogue.has_primary_key(catalogue_key):
                    for col in service_cols:
                        new_value = tuple(set(self._catalogue[catalogue_key][col] + product[col]))
                        if new_value != self._catalogue[catalogue_key][col]:
                            self._catalogue[catalogue_key][col] = new_value
                else:
                    self._catalogue.add_product(product)

                if catalogue_key not in self._catalogue_to_table_row_idx_mapping:
                    self._catalogue_to_table_row_idx_mapping[catalogue_key] = []
                self._catalogue_to_table_row_idx_mapping[catalogue_key].append(self._data_row_indices[0] + idx)

        return self._catalogue

    @property
    def columns(self) -> Tuple[Union[Column, Service]]:
        """Columns in the doc table as Column, Service enums, in order"""
        return COLS_BY_MISSION_AND_GROUP[self._mission][self._group]

    @property
    def col_headers(self) -> Tuple[str, ...]:
        """Column names in the doc table, in order"""
        return tuple(COLUMN_TO_COL_NAME_MAP_BY_GROUP[c][self._group] for c in self.columns)

    @property
    def _data_row_indices(self) -> Tuple[int]:
        """Indices of data (non-header) rows in the table data array"""
        if self.n_data_rows == 0:
            return ()
        return tuple(range(max(self._header_row_indices) + 1, len(self._table.rows)))

    @property
    def n_data_rows(self) -> int:
        """Current number of data (non-header) rows in the table data array"""
        return len(self._table.rows) - len(self._header_row_indices)

    @property
    def data(self) -> Optional[Iterator[dict]]:
        if self._data_row_indices:
            columns = self.columns
            for idx, row in enumerate(self._table.rows[self._data_row_indices[0]: self._data_row_indices[-1] + 1]):
                row_dict = {columns[i]: cell.text.strip() for i, cell in enumerate(row.cells)}
                row_dict.update({'data_row_idx': idx})
                yield row_dict

    def add_product_as_row(
        self,
        product: Union['Product', Iterable['Product']],
        _split_over_two_rows: bool = True,
        _add_to_catalogue: bool = True,
    ) -> None:
        """
        Adds a Product instance, or List/Tuple of Product instances, as row(s)
        in the table. Recurses if List/Tuple of Product instances
        
        Parameters
        ----------
        product: Product, Iterable[Product]
            Product(s) to add to the table

        _split_over_two_rows: bool
            Whether to split the product over two rows e.g. the first for 
            producer entities, the last for consumer entities. Default is True

        _add_to_catalogue: bool
            Whether to add the product to self.catalogue. Default is True
        """
        if isinstance(product, Iterable):
            for product_ in product:
                self.add_product_as_row(
                    product_,
                    _split_over_two_rows=_split_over_two_rows,
                    _add_to_catalogue=_add_to_catalogue,
                )
            return

        if _add_to_catalogue:
            self.catalogue.add_product(product)

        # Producers/consumers to be split over two rows
        if _split_over_two_rows:
            producers_product = deepcopy(product)
            consumers_product = deepcopy(product)

            service_cols = [c for c in self.columns if isinstance(c, Service)]
            for service in service_cols:
                producers_product[service] = tuple(e for e in producers_product[service] if e in PRODUCER_ENTITIES)
                consumers_product[service] = tuple(e for e in consumers_product[service] if e in CONSUMER_ENTITIES)

            self.add_product_as_row(
                [producers_product, consumers_product],
                _split_over_two_rows=False,
                _add_to_catalogue=False,
            )
            return

        # Cache to avoid overheads
        new_row_cells = self._table.add_row().cells
        try:
            row_above_cells = self._table.rows[-2].cells
        except IndexError:
            row_above_cells = None

        # Write each column's values to the new row's appropriate cell
        for idx, column in enumerate(self.columns):
            cell = new_row_cells[idx]
            cell_above = row_above_cells[idx] if row_above_cells else None

            new_value = product[column]
            if new_value is None:
                new_value = ''
            elif not isinstance(new_value, str):
                if isinstance(new_value, (Mission, Group)):
                    new_value = new_value.value
                elif isinstance(new_value, Tuple):
                    # Separate each Entity by a newline character
                    if all([isinstance(item, Entity) for item in new_value]):
                        new_value = '<br>'.join(item.value for item in new_value)
                    else:
                        raise ValueError(f"Can not cast {new_value} to str")
                else:
                    raise TypeError(
                        f"{new_value} must be a str, Mission, Group, or "
                        f"Tuple[Entity], not {type(new_value)}"
                    )

            self.__html_parser.add_html_to_cell(new_value, cell)
            self._set_cell_font_size(cell, self.BODY_FONT_SIZE_PT, bold=False)
            self._set_cell_vertical_alignment(cell, alignment='center')
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if self._word_generator.tracking_changes:
                    self._word_generator.insert_existing_runs(paragraph)

            # Only merge with above if not a Service column
            if not isinstance(column, Service) and cell_above:
                cell_text = ''.join(t.text.strip() or '' for t in cell._tc.iter(qn('w:t')))
                cell_above_text = ''.join(t.text.strip() or '' for t in cell_above._tc.iter(qn('w:t')))
                if cell_text == cell_above_text:
                    # Merge with cell above if sharing the same cell value
                    cell_above.merge(cell)
                    self._set_cell_font_size(cell_above, self.BODY_FONT_SIZE_PT, bold=False)
                    self._set_cell_vertical_alignment(cell_above, alignment='center')

                    for idx, paragraph in enumerate(cell_above.paragraphs):
                        if idx != 0:
                            # Properly clear unnecessary paragraphs from newly merged cell
                            paragraph._element.getparent().remove(paragraph._element)
                        else:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            elif isinstance(column, Service):
                self._set_cell_background(cell, self.COLUMN_COLOURS[column])

    def _set_cell_vertical_alignment(self, cell: Union[docx.table._Cell, CT_Tc], alignment: str = "center"):        
        if not isinstance(cell, (docx.table._Cell, CT_Tc)):
            raise TypeError('Cell must be docx.table._Cell or docx.oxml.table.CT_Tc instance, not {type(cell)}')

        tc = cell if isinstance(cell, CT_Tc) else cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), alignment)
        tcPr.append(vAlign)

    def _set_cell_alignment(self, tc: Union[docx.table._Cell, CT_Tc], alignment: str = 'center'):
        """alignment: 'left', 'center', 'right'"""
        if isinstance(tc, docx.table._Cell):
            tc = tc._tc

        for p in tc.findall(qn('w:p')):
            pPr = p.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p.insert(0, pPr)
            
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), alignment)
            pPr.append(jc)

    def _set_cell_background(self, cell: Union[docx.table._Cell, CT_Tc], fill_hex: str):
        if not isinstance(cell, (docx.table._Cell, CT_Tc)):
            raise TypeError('Cell must be docx.table._Cell or docx.oxml.table.CT_Tc instance, not {type(cell)}')

        tc = cell if isinstance(cell, CT_Tc) else cell._tc
        tcPr = tc.get_or_add_tcPr()

        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex)
        tcPr.append(shd)

    def _set_cell_font_size(self, cell, size_pt: float, bold: bool = False):
        """Sets font size for a specific cell"""
        from docx.shared import Pt
        from lxml import etree

        for paragraph in cell.paragraphs:
            # Cover existing runs
            for run in paragraph.runs:
                run.font.size = Pt(size_pt)
                run.font.bold = bold
            
            # Set default at paragraph level via XML directly
            pPr = paragraph._element.get_or_add_pPr()
            rPr = pPr.find(qn('w:rPr'))
            if rPr is None:
                rPr = etree.SubElement(pPr, qn('w:rPr'))
            
            sz = rPr.find(qn('w:sz'))
            if sz is None:
                sz = etree.SubElement(rPr, qn('w:sz'))
            sz.set(qn('w:val'), str(size_pt * 2))  # half-points

    def _set_table_font_size(
        self,
        body_size_pt: float = BODY_FONT_SIZE_PT,
        header_size_pt: float = HEADER_FONT_SIZE_PT,
        bold_header: bool = BOLD_HEADER
    ):
        """Sets font size across the table"""
        for idx, row in enumerate(self._table.rows):
            is_header_row = idx in self._header_row_indices
            size_pt = header_size_pt if is_header_row else body_size_pt
            for cell in row.cells:
                self._set_cell_font_size(
                    cell, size_pt, bold=is_header_row and bold_header
                )

    def _set_table_borders(self, tbl, color='000000'):
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        tblBorders = OxmlElement('w:tblBorders')
        
        border_config = {
            'top':     {'sz': 12},  # 1.5pt
            'bottom':  {'sz': 12},  # 1.5pt
            'insideH': {'sz': 12},  # 1.5pt
            'left':    {'sz': 8},   # 1pt
            'right':   {'sz': 8},   # 1pt
            'insideV': {'sz': 8},   # 1pt
        }

        for side, config in border_config.items():
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(config['sz']))
            border.set(qn('w:color'), color)
            border.set(qn('w:space'), '0')
            tblBorders.append(border)

        tblPr.append(tblBorders)


@dataclass
class Product:
    """
    Data class representing a product including its producers and consumer 
    (within entities relations). Must have a valid mission, group, and 
    product_type to instantiate.
    
    Parameters
    ----------
    id: str
        Hexadecimal id of the product in the DB e.g. 
        '9df38970_dd91_11ef_80ae_0242c0a83005'

    mission: Mission
        Mission the product belongs to e.g. Mission.S1 (Sentinel-1)

    group: Group
        Group the product belongs to e.g. Group.PRODUCTS (Products)

    product_type: str
        Product type of the Product e.g. 'S[1..6]_RAW__0C'

    payload: str
        Mission payload e.g. 'SAR'

    level: str
        Product processing level e.g. 'L0'

    mode: str
        Product mode e.g. 'SM', 'EW', 'WV'

    type: str
        Product type e.g. 'SLC', 'OCN', 'RAW'

    description: str
        Qualitative description of the product. NOTE: May contain raw HTML

    entities_relations: dict[Service, Tuple[Entity]]
        Dictionary of entity relations whereby keys are services, and values
        are tuples of their producers and consumers. For example:
        
        {
            Service.DA: (Entity.P, Entity.C_PR, Entity.C_LTA),
            Service.LTA: (Entity.P, Entity.C_PR),
            Service.MPC: (Entity.C_DA, ),
            Service.PR: (Entity.P0, )
        }
        
        NOTE: The staticmethod Product.parse_json_service_entities_relation 
        can be used to transform the typical DB entities_relations (list of
        strings) of entities relations to the above format for initialisation

    Notes
    -----
    Access to data attributes via Column or Service instances is possible through
    __getitem__ (see Product.COL_TO_ATTR_MAP and COLS_BY_MISSION_AND_GROUP 
    mappings for Column and Service instances, respectively).
    """
    # Mapping between column headers and JSON fields
    COL_TO_ATTR_MAP: ClassVar[Dict[Column, str]] = {
        Column.DESCRIPTION: 'description',
        Column.PRODUCT_TYPE: 'product_type',
        Column.PAYLOAD: 'payload',
        Column.LEVEL: 'level',
        Column.MODE: 'mode',
        Column.TYPE: 'type',
    }

    __slots__ = (
        'id', 'mission', 'group', 'product_type',
        'payload', 'level', 'mode', 'type',
        'description', 'entities_relations',
    )

    id: str
    mission: Mission
    group: Group
    product_type: str
    payload: str
    level: str
    mode: str
    type: str
    description: str
    entities_relations: dict[Service, Tuple[Entity]]

    def __post_init__(self):
        if not self.product_type:
            raise ValueError("Product requires a product_type/name")
        if not self.mission:
            raise ValueError("Product requires a mission")
        if not self.group:
            raise ValueError("Product requires a group")

    def __getitem__(self, key: Union[Column, Service]) -> Union[Mission, Group, str, Tuple[Entity]]:
        if isinstance(key, Column):
            return self.__getattribute__(self.COL_TO_ATTR_MAP[key])
        elif isinstance(key, Service):
            try:
                return self.entities_relations[key]
            except KeyError as err:
                if key in COLS_BY_MISSION_AND_GROUP[self.mission][self.group]:
                    return ()
                raise err
        else:
            raise TypeError(f"key must be a Column or Service instance, not {type(key)}")

    def __setitem__(self, key: Union[Column, Service], value: Union[str, Mission, Group, Service, Tuple[Entity]]):
        self[key]  # Force checks in __getitem__ to run
        if isinstance(key, Service):
            self.entities_relations[key] = value
        elif isinstance(key, Column):
            self.__setattr__(self.COL_TO_ATTR_MAP[key], value)

    @classmethod
    def from_json(cls, json_obj: Dict[str, Union[str, List[str]]]) -> 'Product':
        """Create a Product instance from a JSON object returned by /rest/api/dataflow/<config_id> endpoint"""
        for mission in Mission:
            if mission.name == json_obj["mission"].strip():
                break
        else:
            mission = json_obj["mission"].strip()
            raise ValueError(f"Mission '{mission}' not found")

        entities_relations = {}
        for relation in json_obj['entities_relations']:
            service, entities = cls.parse_json_service_entities_relation(relation)
            entities_relations[service] = entities

        instance = cls(
            id=json_obj["id"].strip(),
            mission=mission,
            group=GROUP_NAME_DATAFLOW_MAP[json_obj["group"]],
            product_type=json_obj["name"].strip() if json_obj["name"] else None,
            payload=json_obj["payload"].strip() if json_obj["payload"] else None,
            level=json_obj["level"].strip() if json_obj["level"] else None,
            mode=json_obj["sensor-mode"].strip() if json_obj["sensor-mode"] else None,
            type=json_obj["type"].strip() if json_obj["type"] else None,
            description=json_obj["description"].strip() if json_obj["description"] else None,
            entities_relations=entities_relations
        )

        return instance
    
    @staticmethod
    def parse_json_service_entities_relation(
        service_entity_relation: str,
        relation_delim: str = ':',
        entity_delim: str = ';',
        consumer_delim: str = '/',
    ) -> Tuple[Service, Tuple[Entity]]:
        """
        Parses the 'entities_relations' field from the JSON data, which is a 
        list of strings in the format "Service: relation1; relation2; ...".

        Parameters
        ----------
        service_entity_relation : list[str]
            A string containing a service and its related entities of the format
            "<service>: <entity1>; <consumer_entity1>/<consumer_entity2>; ..."
            where ':' is the default value for the relation_delim, ';' is the
            default value for the entity_delim and '/' is the default value for
            the consumer_delim.
        relation_delim : str, optional
            The delimiter between the service and the related entities in the
            input string (default is ':').
        entity_delim : str, optional
            The delimiter between related entities in the input string
            (default is ';').
        consumer_delim : str, optional
            The delimiter between consumer entities in the input string
            (default is '/').

        Returns
        -------
        Tuple[Service, Tuple[Entity]]
            A tuple containing a Service and a tuple of its related entities.
        """
        service_str, entities_str = [_.strip() for _ in service_entity_relation.split(relation_delim)]

        service = Service(service_str)
        entities = []
        for string in entities_str.split(entity_delim):
            for entity_str in string.split(consumer_delim):
                entity_str = entity_str.strip()
                try:
                    entities.append(Entity(entity_str))
                except ValueError as err:
                    if entity_str.startswith('-'):
                        entities.append(Entity(f'C{entity_str}'))
                    else:
                        raise err

        return service, tuple(entities)


# Types related to ProductCatalogue class
CataloguePrimaryKey: TypeAlias = Tuple[Union[Column, Service, str], ...]
CataloguePrimaryKeyFunc: TypeAlias = Callable[[Product], CataloguePrimaryKey]

class ProductCatalogue:
    """
    Queryable catalogue of products
    
    Parameters
    ----------
    products: Iterable[Product]
        Products to be added to the catalogue at initialisation (optional)
    
    pk_func: CataloguePrimaryKeyFunc
        'Primary Key Function'. Callable whose single argument is a Product 
        instance, and returns a unique CataloguePrimaryKey to be used to 
        identify the product in the catalogue. Default is to return a tuple of
        the products mission, group, and product type values.

    Attributes
    ----------
    _products: List[Product]
        List of products added to the catalogue. Read-only access available through products property
    
    _catalogue: Dict[Union[Column, Service], Product]
        Underlying dictionary containing the primary keys and products as key/value pairs, respectively

    _pk_func: CataloguePrimaryKeyFunc
        Primary key function used to determine unique keys for each added product        
    """
    _DEFAULT_PK = lambda product: (product.mission, product.group, product[Column.PRODUCT_TYPE])

    def __init__(self, products: Iterable[Product] = None, pk_func: CataloguePrimaryKeyFunc = _DEFAULT_PK):
        """
        Initialise the catalogue. Loads any supplied Product instances
        (products arg) into the catalogue
        """
        self._products: List[Product] = []
        self._catalogue: Dict[Union[Column, Service], Product] = {}
        self._pk_func: CataloguePrimaryKeyFunc = pk_func
        self.track_changes: bool = True

        if products is not None:
            for product in products:
                self.add_product(product)

    def __getitem__(self, key: Tuple[Union[str, Column, Service], ...]) -> Product:
        """Retrieves product by primary key"""
        return self._catalogue[key]

    def __repr__(self):
        return f"{self.__class__.__name__} [{len(self._products)} products across {len(set((_.mission, _.group) for _ in self._products))} mission-groups]"

    def __iter__(self) -> Iterator[int]:
        return iter(self.keys)

    @property
    def empty(self) -> bool:
        """
        Returns False if the dictionary has added products, True otherwise
        """
        return len(self._products) == 0

    @property
    def products(self) -> List[Product]:
        """Returns list of Product instances added to the catalogue"""
        return self._products

    @property
    def keys(self) -> Tuple[Tuple[Union[str, Column, Service], ...], ...]:
        """Returns tuple of primary keys in the catalogue"""
        return tuple(self._catalogue.keys())

    def add_product(self, product: Product):
        """
        Adds product to catalogue
        
        Parameters
        ----------
        product: Product
            Product to add to catalogue
        
        Raises
        ------
        KeyError:
            In cases where the product-to-be-added's evaluated key (using 
            self._pk_func) is already present in the catalogue
        """
        key = self._pk_func(product)
        if key in self._catalogue:
            raise KeyError(f"Product with key '{key}', already in catalogue")

        self._products.append(product)
        self._catalogue[key] = product

    def get_group_products(self, group: Group) -> 'ProductCatalogue':
        """
        Creates a new ProductCatalogue containing only products with 
        group == group
        """
        return ProductCatalogue(filter(lambda product: product.group == group, self.products), self._pk_func)

    def get_mission_and_group_products(self, mission: Mission, group: Group) -> 'ProductCatalogue':
        """
        Creates a new ProductCatalogue containing only products with 
        group == group and mission == mission
        """
        return ProductCatalogue(filter(lambda product: product.mission == mission and product.group == group, self.products), self._pk_func)

    def get_mission_products(self, mission: Mission) -> 'ProductCatalogue':
        """
        Creates a new ProductCatalogue containing only products with 
        mission == mission
        """
        return ProductCatalogue(filter(lambda product: product.mission == mission, self.products), self._pk_func)

    def has_primary_key(self, pk: Tuple[Union[str, Column, Service], ...]) -> bool:
        """Checks if the catalogue contains the key 'pk'"""
        return self._catalogue.get(pk, None) is not None

    def search_by_column(self, search_term: re.Pattern, column: Union[Column, Entity, str]) -> List[Product]:
        """
        Search a specific column's values by regex pattern and return matches
        as a list of Products
        """
        return list(filter(lambda product: search_term.search(product[column]), self.products))

class DataflowDocCreator:
    """
    Singleton responsible for dataflow document creation across the app
    """
    _instance = None
    _instance_lock = threading.Lock()

    def __new__(cls, *args, **kwargs) -> 'DataflowDocCreator':
        with cls._instance_lock:
            if cls._instance is None:
                cls._instance = super().__new__(cls)
                cls._instance._initialised = False
        return cls._instance

    def __init__(self, app: flask.app.Flask):
        if self._initialised:
            return

        self._app = app
        self._lock = threading.Lock()
        self._stop_event = threading.Event()
        self._thread = None
        self._initialised = True

        self.track_changes = True
        self.config_id = None

    @property
    def logger(self) -> logging.Logger:
        return self._app.logger

    def trigger(self):
        if self.config_id is None:
            raise ValueError("Set config_id before running")

        with self._lock:
            # Cancel any running job
            if self._thread and self._thread.is_alive():
                self.logger.info("Interrupting existing document creation job")
                self._stop_event.set()
                self._thread.join()  # wait for clean exit

            # Reset and launch fresh
            self._stop_event.clear()
            self._thread = threading.Thread(
                target=self._run,
                daemon=True
            )
            self._thread.start()

    def _run(self):
        try:
            unofficial_doc_buffer = self._create_document(official=False)
            official_doc_buffer = self._create_document(official=True)

            if self._stop_event.is_set():
                return

        except Exception as err:
            self.logger.exception(f"Document creation failed: {err}")

        self._upload_to_s3(unofficial_doc_buffer, official=False)
        self._upload_to_s3(official_doc_buffer, official=True)

    def _create_document(self, official: bool = False) -> io.BytesIO:
        from apps.models.nosql.Graph import Graph

        graph = Graph()
        scen_graph = graph.find({'id': self.config_id})
        json_objs = json.loads(scen_graph[0]['graph'])['product_types']

        json_catalogue = load_dataflow_json_to_catalogue(json_objs)
        template = WordGenerator(CURRENT_DATAFLOW_DOC_TEMPLATE)
        template.tracking_changes = self.track_changes

        # Write in json_catalogue to empty template doc (which already has table headers)
        for mission in Mission:
            for group in Group:
                # If stop event has been triggered, terminate the document creation
                if self._stop_event.is_set():
                    return

                db_mission_group_products = json_catalogue.get_mission_and_group_products(mission, group)
                if len(db_mission_group_products.products) == 0:
                    continue

                mission_group_doc_tables = template.get_mission_group_table(mission, group)
                if len(mission_group_doc_tables) == 0:
                    continue

                mission_group_doc_table = DataflowDocTable(template, mission, group)
                products = db_mission_group_products.products

                def sort_key(prod: Product) -> str:
                    key = ""
                    for col in mission_group_doc_table.columns:
                        val = prod[col]

                        if issubclass(val.__class__, Enum):
                            val = val.value

                        key += str(val)

                    return key

                products.sort(key=sort_key)
                for product in db_mission_group_products.products:
                    mission_group_doc_table.add_product_as_row(product)

        buffer = io.BytesIO()
        template.document.save(buffer)
        buffer.seek(0)

        return buffer

    def _upload_to_s3(self, doc_buffer: io.BytesIO, official: bool = False):
        from apps import s3_client

        s3_obj_key = self._app.config['S3_DATAFLOW_DOC_OFFICIAL_KEY'] if official else self._app.config['S3_DATAFLOW_DOC_UNOFFICIAL_KEY']
        s3_bucket = self._app.config['S3_BUCKET']
        doc_type_str = "official" if official else "unofficial"

        try:
            s3_client.put_object(
                Body=doc_buffer,
                Bucket=s3_bucket,
                Key=s3_obj_key
            )

            
            self.logger.info(f"Successfully uploaded {doc_type_str} doc to {s3_obj_key} in {s3_bucket} bucket")

        except Exception as err:
            self.logger.error(f"Could not upload {doc_type_str} doc to {s3_obj_key} in {s3_bucket} bucket: {err}")

# endregion
# ########################################################################### #
# ############################ Mappings ##################################### #
# ########################################################################### #
# region mappings
# Which data columns are present in each table, by mission and group, in 
# order
COLS_BY_MISSION_AND_GROUP: Dict[Mission, Dict[Group, Tuple[Union[Column, Service], ...]]] = {
    Mission.S1: {
        Group.PRODUCTS: (
            Column.PAYLOAD, Column.MODE, Column.LEVEL, Column.TYPE, Column.PRODUCT_TYPE,
            Service.PR, Service.FOS, Service.MPC, Service.LTA, Service.DA, Service.POD
        ),
        Group.AUX_DATA: (
            Column.PRODUCT_TYPE, Column.DESCRIPTION,
            Service.FOS, Service.MP, Service.ADG, Service.POD, Service.MPC, Service.PR, Service.LTA, Service.DA, Service.X_BAND
        ),
        Group.MP_AND_FOS: (
            Column.PRODUCT_TYPE, Column.DESCRIPTION,
            Service.FOS, Service.MP, Service.ADG, Service.POD, Service.MPC, Service.PR, Service.EDRS, Service.DA, Service.X_BAND
        ),
        Group.OLQC_REPORTS: (
            Column.PRODUCT_TYPE,
            Service.PR, Service.E2E, Service.MPC, Service.LTA, Service.DA
        ),
        Group.REMOVED_PRODUCTS: (
            Column.PAYLOAD, Column.MODE, Column.LEVEL, Column.TYPE, Column.PRODUCT_TYPE,
            Service.PR, Service.FOS, Service.MPC, Service.LTA, Service.DA, Service.POD
        ),
        Group.REMOVED_AUX_DATA: (
            Column.PRODUCT_TYPE, Column.DESCRIPTION,
            Service.FOS, Service.MP, Service.ADG, Service.POD, Service.MPC, Service.PR, Service.LTA, Service.DA, Service.X_BAND
        ),
        Group.REMOVED_MP_AND_FOS: (
            Column.PRODUCT_TYPE, Column.DESCRIPTION,
            Service.FOS, Service.MP, Service.ADG, Service.POD, Service.MPC, Service.PR, Service.EDRS, Service.DA, Service.X_BAND
        ),
    },
    Mission.S2: {
        Group.PRODUCTS: (
            Column.PAYLOAD, Column.LEVEL, Column.TYPE, Column.PRODUCT_TYPE,
            Service.POD, Service.PR, Service.FOS, Service.MPC, Service.LTA, Service.DA
        ),
        Group.AUX_DATA: (
            Column.PRODUCT_TYPE, Column.DESCRIPTION,
            Service.FOS, Service.MP, Service.ADG, Service.POD, Service.MPC, Service.PR, Service.LTA, Service.DA
        ),
        Group.MP_AND_FOS: (
            Column.PRODUCT_TYPE, Column.DESCRIPTION,
            Service.FOS, Service.MP, Service.ADG, Service.POD, Service.MPC, Service.PR, Service.DA
        ),
        Group.OLQC_REPORTS: (
            Column.PRODUCT_TYPE,
            Service.PR, Service.E2E, Service.MPC, Service.LTA, Service.DA
        ),
    },
    Mission.S3: {
        Group.PRODUCTS: (
            Column.PAYLOAD, Column.LEVEL, Column.DESCRIPTION, Column.PRODUCT_TYPE,
            Service.PR, Service.MPC, Service.LTA, Service.DA, Service.EUM, Service.POD
        ),
        Group.AUX_DATA: (
            Column.PRODUCT_TYPE, Column.DESCRIPTION,
            Service.EXT, Service.EUM, Service.ADG, Service.POD, Service.RS, Service.PR, Service.LTA, Service.DA, Service.MPC
        ),
        Group.REMOVED_PRODUCTS: (
            Column.PAYLOAD, Column.LEVEL, Column.DESCRIPTION, Column.PRODUCT_TYPE,
            Service.PR, Service.MPC, Service.LTA, Service.DA, Service.EUM, Service.POD
        ),
        Group.REMOVED_AUX_DATA: (
            Column.PRODUCT_TYPE, Column.DESCRIPTION,
            Service.EXT, Service.EUM, Service.ADG, Service.POD, Service.RS, Service.PR, Service.LTA, Service.DA, Service.MPC
        ),
    },
    Mission.S5P: {
        Group.PRODUCTS: (
            Column.PAYLOAD, Column.LEVEL, Column.PRODUCT_TYPE,
            Service.MPC, Service.PR, Service.LTA, Service.DA
        ),
        Group.AUX_DATA: (
            Column.PRODUCT_TYPE, Column.DESCRIPTION,
            Service.FOS, Service.MPC, Service.PR, Service.LTA, Service.DA
        ),
        Group.REMOVED_PRODUCTS: (
            Column.PAYLOAD, Column.LEVEL, Column.PRODUCT_TYPE,
            Service.MPC, Service.PR, Service.LTA, Service.DA
        ),
        Group.REMOVED_AUX_DATA: (
            Column.PRODUCT_TYPE, Column.DESCRIPTION,
            Service.FOS, Service.MPC, Service.PR, Service.LTA, Service.DA
        ),
    },
    Mission.S6: {
        Group.AUX_DATA: (
            Column.PRODUCT_TYPE, Column.DESCRIPTION,
            Service.POD, Service.DA
        ),
    }
}

# Mapping between Column (and Service acting as columns) enum and column 
# header text
COLUMN_TO_COL_NAME_MAP_BY_GROUP: Dict[Union[Column, Service], Dict[Group, str]] = {
    **{c: {g: c.value for g in Group} for c in Column},
    **{
        Service.ADG: {g: "ADG" for g in Group},
        Service.DA: {
            g: "DATA ACQUISITION" if g in (Group.MP_AND_FOS, Group.REMOVED_MP_AND_FOS)
                else "DA"
                for g in Group
        },
        Service.E2E: {
            g: "E2EM" if g in (Group.OLQC_REPORTS, )
                else "E2E"
                for g in Group
        },
        Service.EDRS: {g: "EDRS" for g in Group},
        Service.EUM: {g: "EUM" for g in Group},
        Service.EXT: {g: "EXT Provider" for g in Group},
        Service.FOS: {g: "FOS" for g in Group},
        Service.LTA: {g: "LTA" for g in Group},
        Service.MP: {g: "MP" for g in Group},
        Service.MPC: {g: "MPC" for g in Group},
        Service.POD: {g: "POD" for g in Group},
        Service.PR: {g: "PR" for g in Group},
        Service.RS: {g: "RS" for g in Group},
        Service.X_BAND: {g: "Col X-Band" for g in Group},
    }
}


# Map between Group enum and the corresponding group name in the dataflow data
DATAFLOW_GROUP_NAME_MAP = {
    Group.PRODUCTS: "Products",
    Group.AUX_DATA: "AUX Data",
    Group.MP_AND_FOS: "MP and FOS files",
    Group.OLQC_REPORTS: "OLQC Reports",
    Group.REMOVED_PRODUCTS: "Removed Products",
    Group.REMOVED_AUX_DATA: "Removed AUX Data",
    Group.REMOVED_MP_AND_FOS: "Removed MP and FOS files"
}
GROUP_NAME_DATAFLOW_MAP = {v: k for k, v in DATAFLOW_GROUP_NAME_MAP.items()}

# Mapping of DB payload name to its name in the document
DB_TO_DOC_PAYLOAD_NAME_MAP = {
    'SAR': 'SAR Instrument',
    'AIS': 'AIS instrument',
    'MSI': 'Multi-Spectral Instrument (MSI)',
    'HKTM': 'Telemetry Data',
    'SYN': 'Synergy (OLCI + SLSTR)',
}

# endregion
# ########################################################################### #
# ############################ Functions #################################### #
# ########################################################################### #
# region functions

def load_dataflow_json_to_catalogue(json_obj: List[Dict], logger: Optional[logging.Logger] = None) -> ProductCatalogue:
    if logger is None:
        logger = logging.getLogger(__name__)

    json_catalogue = ProductCatalogue()

    for json_entry in json_obj:
        try:
            json_catalogue.add_product(Product.from_json(json_entry))
        except Exception as err:
            logger.warning(f"Can not create Product from {json.dumps(json_entry)} due to: {err}")
    
    return json_catalogue


def write_products_to_empty_dataflow_doc(json_objs: List[Dict], track_changes: bool = True) -> WordGenerator:
    json_catalogue = load_dataflow_json_to_catalogue(json_objs)
    template = WordGenerator(CURRENT_DATAFLOW_DOC_TEMPLATE)
    template.tracking_changes = track_changes

    # Write in json_catalogue to empty template doc (which already has table headers)
    for mission in Mission:
        for group in Group:
            db_mission_group_products = json_catalogue.get_mission_and_group_products(mission, group)
            if len(db_mission_group_products.products) == 0:
                continue

            mission_group_doc_tables = template.get_mission_group_table(mission, group)
            if len(mission_group_doc_tables) == 0:
                continue

            mission_group_doc_table = DataflowDocTable(template, mission, group)
            products = db_mission_group_products.products

            def sort_key(prod: Product) -> str:
                key = ""
                for col in mission_group_doc_table.columns:
                    val = prod[col]

                    if issubclass(val.__class__, Enum):
                        val = val.value

                    key += str(val)

                return key

            products.sort(key=sort_key)
            for product in db_mission_group_products.products:
                mission_group_doc_table.add_product_as_row(product)

    return template
# endregion

if __name__ == '__main__':
    import logging
    import shutil

    logging.basicConfig(stream=sys.stdout, level=logging.DEBUG, format='%(levelname)s: %(message)s')
    logger = logging.getLogger(__name__)
    # test = WordGenerator("/mnt/c/Users/simon/OneDrive/S3TeamKnowledgebase/OCS/docs/[EOF-DFC] EOF_CSC Sentinels Data Flow Configuration 1.7.docx")
    def main():
        json_path = Path('/data/ocs/Configuration-Tool/apps/utils/assets/dataflow_data.json')
        with json_path.open(encoding='utf-8') as fd:
            json_obj = json.load(fd)

        doc = write_products_to_empty_dataflow_doc(json_obj)

        save_file = doc.save('temp')
        shutil.move(save_file, "/mnt/c/Users/simon/Desktop/test.docx")
  
  
    # import cProfile
    # cProfile.run('main()?')
    # import line_profiler
    # profiler = line_profiler.LineProfiler()
    # profiler.add_function(DataflowDocTable.add_product_as_row)  # pass the function directly
    # profiler.add_function(ProductCatalogue.get_mission_and_group_products)

    # profiler.enable()
    main()  # your entry point
    # profiler.disable()
    pass
    # raise SystemExit(main())
